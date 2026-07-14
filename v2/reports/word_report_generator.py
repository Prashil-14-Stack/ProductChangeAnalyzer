"""
==========================================================
Word Report Generator

Purpose
-------
Generates a professional Microsoft Word report from a
ComparisonResult.

Responsibilities
----------------
✓ Executive Summary
✓ Comparison Statistics
✓ High Impact Changes
✓ Detailed Comparison Table

Future
------
✓ Company Branding
✓ Logo
✓ Table of Contents
✓ Hyperlinks
✓ Revision History

==========================================================
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class WordReportGenerator:

    """
    Generates Word reports from ComparisonResult.
    """

    def __init__(self, output_folder="output"):

        self.output_folder = Path(output_folder)

        self.output_folder.mkdir(
            parents=True,
            exist_ok=True
        )

    # ======================================================
    # Public
    # ======================================================

    def generate(
        self,
        comparison_result,
        filename="Product_Comparison_Report.docx"
    ):

        document = Document()

        self._add_title(
            document,
            comparison_result
        )

        self._add_summary(
            document,
            comparison_result
        )

        self._add_high_impact_changes(
            document,
            comparison_result
        )

        self._add_detailed_comparison(
            document,
            comparison_result
        )

        output_path = self.output_folder / filename

        document.save(output_path)

        return output_path

    # ======================================================
    # Private
    # ======================================================

    def _add_title(
        self,
        document,
        comparison_result
    ):

        heading = document.add_heading(
            "Product Comparison Report",
            level=1
        )

        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        paragraph = document.add_paragraph()

        paragraph.add_run(
            "Product V1: "
        ).bold = True

        paragraph.add_run(
            comparison_result.product_name_v1
        )

        paragraph.add_run("\n")

        paragraph.add_run(
            "Product V2: "
        ).bold = True

        paragraph.add_run(
            comparison_result.product_name_v2
        )

    # ------------------------------------------------------

    def _add_summary(
        self,
        document,
        comparison_result
    ):

        document.add_heading(
            "Executive Summary",
            level=2
        )

        table = document.add_table(
            rows=1,
            cols=2
        )

        table.style = "Table Grid"

        header = table.rows[0].cells

        header[0].text = "Metric"

        header[1].text = "Value"

        for key, value in comparison_result.summary.items():

            row = table.add_row().cells

            row[0].text = str(key)

            row[1].text = str(value)

    # ------------------------------------------------------

    def _add_high_impact_changes(
        self,
        document,
        comparison_result
    ):

        document.add_heading(
            "High Impact Changes",
            level=2
        )

        for item in comparison_result.high_impact_items():

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.add_run(
                item.parameter_name
            ).bold = True

            paragraph.add_run(
                f" ({item.status})"
            )

    # ------------------------------------------------------

    def _add_detailed_comparison(
        self,
        document,
        comparison_result
    ):

        document.add_heading(
            "Detailed Comparison",
            level=2
        )

        table = document.add_table(
            rows=1,
            cols=5
        )

        table.style = "Table Grid"

        headers = table.rows[0].cells

        headers[0].text = "Parameter"

        headers[1].text = "Status"

        headers[2].text = "Impact"

        headers[3].text = "Old Value"

        headers[4].text = "New Value"

        for item in comparison_result.items:

            row = table.add_row().cells

            row[0].text = item.parameter_name

            row[1].text = item.status

            row[2].text = item.impact

            row[3].text = str(item.old_value)

            row[4].text = str(item.new_value)