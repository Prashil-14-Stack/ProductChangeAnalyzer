from docx import Document


def extract_cell_content(cell):

    content = []

    if cell.text.strip():
        content.append(cell.text.strip())

    for table in cell.tables:

        for row in table.rows:

            row_data = []

            for nested_cell in row.cells:

                value = nested_cell.text.strip()

                if value:
                    row_data.append(value)

            if row_data:
                content.append(
                    " | ".join(row_data)
                )

    return "\n".join(content)


class WordReader:

    def read(self, file_path):

        document = Document(file_path)

        data = {}

        for table in document.tables:

            for row in table.rows:

                if len(row.cells) >= 2:

                    parameter = row.cells[0].text.strip()

                    description_parts = []

                    for cell in row.cells[1:]:

                        text = extract_cell_content(cell)

                        if text:
                            description_parts.append(text)

                    description = "\n".join(
                        description_parts
                    )

                    if parameter:

                        if parameter in data:

                            data[parameter] += (
                                "\n" + description
                            )

                        else:

                            data[parameter] = description

        return data